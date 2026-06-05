# -*- coding: utf-8 -*-
import copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = '/root/.claude/uploads/fcb12aea-2856-4410-a156-650bc3f4f673/64d27673-__________.docx'
OUT = '/home/user/thesis-formatter/docs/呼伦贝尔学院本科论文规范交付智能助手研究-申报书初稿.docx'

doc = Document(SRC)
tbl = doc.tables[0]
cells = [row.cells[0] for row in tbl.rows]

# ---------- helpers ----------
def style_run(run, size=12, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = 'Times New Roman'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    if color is not None:
        run.font.color.rgb = color

def clear_cell_keep_heading(cell):
    paras = cell.paragraphs
    for p in paras[1:]:
        p._element.getparent().remove(p._element)
    # restyle heading as bold section title
    h = cell.paragraphs[0]
    for r in h.runs:
        style_run(r, size=12, bold=True)

def add_para(cell, text='', *, bold=False, indent=False, size=12, align=None,
             color=None, space_before=2, space_after=2, line=1.5):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    if indent:
        pf.first_line_indent = Pt(size*2)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        style_run(r, size=size, bold=bold, color=color)
    return p

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), '000000')
        borders.append(e)
    tblPr.append(borders)

def fill_cell(cell, text, *, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    style_run(r, size=size, bold=bold)

def add_table_block(cell, caption, headers, rows, note=None):
    cap = add_para(cell, caption, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                   size=10.5, space_before=8, space_after=2)
    t = cell.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = True
    set_table_borders(t)
    for j, h in enumerate(headers):
        fill_cell(t.rows[0].cells[j], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, rowvals in enumerate(rows, start=1):
        for j, v in enumerate(rowvals):
            fill_cell(t.rows[i].cells[j], v)
    if note:
        add_para(cell, note, size=9, color=RGBColor(0x59, 0x59, 0x59),
                 space_before=2, space_after=6)
    else:
        add_para(cell, '', space_before=0, space_after=4)  # trailing para after table

def add_figure_placeholder(cell, figtitle, note, lines=4):
    box = cell.add_table(rows=1, cols=1)
    set_table_borders(box)
    bc = box.rows[0].cells[0]
    bc.text = ''
    ph = bc.paragraphs[0]
    ph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ph.add_run('（此处插入图片，由申报人替换）')
    style_run(r, size=10.5, color=RGBColor(0x88, 0x88, 0x88))
    for _ in range(lines):
        ep = bc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(cell, figtitle, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             size=10.5, space_before=4, space_after=1)
    add_para(cell, note, size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x59, 0x59, 0x59), space_before=0, space_after=8)

L = WD_ALIGN_PARAGRAPH.LEFT
C = WD_ALIGN_PARAGRAPH.CENTER

# ================= ROW 0  选题编号及名称 =================
c = cells[0]; clear_cell_keep_heading(c)
add_para(c, '选题编号：【填写：对应申报指南中的选题编号】', indent=False)
add_para(c, '选题名称（研究方向）：人工智能赋能本科毕业论文质量保障与过程管理'
            '（教育与人才培养研究／数字化与创新研究方向）', indent=False)

# ================= ROW 1  申报题目 =================
c = cells[1]; clear_cell_keep_heading(c)
add_para(c, '呼伦贝尔学院本科论文规范交付智能助手研究', bold=True, align=C, size=14)
add_para(c, '项目类别：人文社会科学类（教育与人才培养研究／数字化与创新研究）')
add_para(c, '研究周期：2026年6月—2028年6月（两年）')
add_para(c, '说明：本项目以呼伦贝尔学院本科毕业论文的“提交—指导—交付”场景为研究对象，依托一套已实际部署并阶段性试运行的论文辅助系统，研究“提交前体检—学术表达优化—格式对齐—标准 Word 导出”的论文规范交付闭环机制，及其校本化、规则化、可评价的实现路径。项目不以软件功能堆叠为目标，而以“将论文指导经验转化为可计算规则、将论文质量问题转化为可诊断指标、将学生修改过程转化为人机协同流程”为研究核心。', indent=True)

# ================= ROW 2  一、国内外研究现状及趋势 =================
c = cells[2]; clear_cell_keep_heading(c)
J = WD_ALIGN_PARAGRAPH.JUSTIFY
add_para(c, '近年来，教育数字化转型已上升为国家战略，被视为促进公平、提升质量、推动高等教育高质量发展的关键路径[1-2]；其本质是技术、业务与人本层面的系统性变革，将本科人才培养的质量保障提升到新的高度，但在落地中仍面临理念重塑与路径选择的现实困境[3]。', indent=True, align=J)
add_para(c, '在这一进程中，以 ChatGPT 为代表的生成式人工智能迅速进入教育场景，被认为将深刻影响教育形态与教学模式[4-5]，对高等教育的价值与使命提出新的追问[6]，并推动以知识点为核心的教学方式变革[7]。值得注意的是，生成式人工智能的能力最直接地作用于语言与文本，学术写作由此成为其落地的关键界面。', indent=True, align=J)
add_para(c, '沿着这一界面，国际研究已就人工智能辅助写作积累了较系统的实证证据：自动化写作评价（AWE）与即时反馈可帮助学习者改进词汇、语法与篇章准确性[8-9]，多项元分析表明其对写作质量具有正向作用[10-11]；近两年的实验进一步显示，大语言模型生成的反馈在促进学生修改、提升写作表现方面可与教师反馈相当甚至更优[12-13]，其在语法纠错与文本润色等任务上亦展现出较强能力[14-15]。这说明，将人工智能用于表达层面的优化在技术上已具备可行性。', indent=True, align=J)
add_para(c, '然而，技术可行并不等于应用恰当。生成式人工智能进入写作过程，同时引发学术诚信、规范与学生主体性等治理议题[16-17]，学界就其基本争议、风险与边界展开讨论[18-19]，并在《中华人民共和国学位法》等制度框架下关注学位与毕业论文写作中的合规使用[20]。由此，“人机协同”逐渐成为被广泛接受的框架——强调以学生为写作主体、以系统为辅助，并对二者贡献的边界进行重新界定[21-22]。', indent=True, align=J)
add_para(c, '进一步聚焦可以发现，上述研究多面向通用写作或研究生科研场景，而落到本科毕业论文这一具体环节，问题更为现实。本科毕业论文是本科人才培养的重要环节，其质量保障正从结果检查转向全过程管理[23]；但抽检与教学实践表明，仍有相当比例的论文存在质量问题，学生学术写作能力与规范意识不足是重要诱因[24-25]，而规范要求长期依赖教师个体经验与口头反馈传递，缺少可计算、可检查、可反馈的工具支撑[26]。', indent=True, align=J)
add_para(c, '与此相应，现有智能写作与文档检查工具多面向通用学术规范，存在校本规范适配不足、重文本生成而轻规范交付、重单点功能而轻提交前闭环、且普遍缺少校内试运行与应用验证等局限；即便有研究将生成式人工智能视为高等教育的范式性变化，也普遍提示其成效需在真实场景中审慎验证[27-28]。对于以服务地方、培养应用型人才为定位的地方本科高校而言，这一缺口尤为突出。综合来看，面向本科毕业论文规范交付这一具体场景，以校本规范为依据、以“提交前体检—学术表达优化—格式对齐—标准导出”闭环为研究对象、以人机协同为机制，并依托已实际运行系统开展数据验证的研究，仍存在明显空白；本项目正是面向这一空白展开（现有工具与本项目定位对比见表0）。', indent=True, align=J)
add_table_block(c, '表0　现有工具与本项目定位对比表',
    ['维度', '现有通用工具的典型特征', '存在的不足', '本项目定位'],
    [['规范适配', '面向通用学术规范', '校本规范适配不足', '依据呼伦贝尔学院本科论文格式规范'],
     ['功能侧重', '重文本生成、轻规范交付', '不解决能否合规提交', '面向规范交付的全流程'],
     ['流程完整性', '重单点功能、轻提交前闭环', '缺少先体检后修改闭环', '体检—优化—对齐—导出闭环'],
     ['知识转化', '重写作辅助、轻规则转化', '校规、经验、错误未打通', '三源融合的校本规则库'],
     ['应用验证', '缺系统试运行与校内数据', '难以评估真实成效', '已有阶段性试运行与反馈数据']])

# ================= ROW 3  二、研究目的与意义 =================
c = cells[3]; clear_cell_keep_heading(c)
add_para(c, '研究目的：围绕呼伦贝尔学院本科论文规范交付场景，研究并完善一套“校本规则库 + 提交前体检 + 学术表达优化 + 格式对齐 + 标准导出”的智能助手，将论文指导中的隐性规范经验显性化、规则化、流程化，并以校内试运行数据验证其应用成效。', indent=True)
add_para(c, '1. 理论意义', bold=True)
add_para(c, '（1）探索本科论文规范交付的规则化、流程化与智能化机制，回答“校本规范知识如何被建模为可计算规则”这一问题；（2）丰富人工智能赋能地方高校人才培养质量保障的应用研究，提供“经验—规则—流程—反馈”的转化范式；（3）为“人机协同修改”在本科论文场景下的合理边界（以学生为写作主体、系统为辅助）提供实证参考。', indent=True)
add_para(c, '2. 实践意义', bold=True)
add_para(c, '（1）降低学生论文提交前盲目修改、反复提交的成本；（2）减少教师在格式与基础表达问题上的重复性审核工作，使指导回归选题、论证与学术内容；（3）提升本科论文的提交质量与一次性合规率；（4）形成呼伦贝尔学院校本化的本科论文规范交付流程与配套工具。', indent=True)
add_para(c, '3. 推广意义', bold=True)
add_para(c, '（1）为地方高校构建低成本、可复制、可迭代的论文质量支持工具提供参考；（2）系统与规则框架可向课程论文、毕业设计、开题报告等教学过程管理环节拓展，具备可持续延伸价值。', indent=True)

# ================= ROW 4  三、研究内容、工作方案及重难点 =================
c = cells[4]; clear_cell_keep_heading(c)
add_para(c, '本项目设置四项研究内容，构成“需求—规则—系统—验证”的研究链条（见图8、图3、图7）。', indent=True)
add_figure_placeholder(c, '图8　研究内容分解图', '图注：四项研究内容构成“需求—规则—系统—验证”的研究链条。')
add_para(c, '研究内容一：呼伦贝尔学院本科论文规范交付需求分析', bold=True)
add_para(c, '梳理学校本科论文格式规范文件，调研指导教师审核关注点与学生常见问题，明确“提交前应检查什么、规范交付包含哪些环节、人机如何分工”，形成需求分析报告与问题清单。', indent=True)
add_para(c, '研究内容二：校本论文规范规则库与质量诊断指标构建', bold=True)
add_para(c, '将学校规范条款、教师审核经验、学生常见错误进行拆解，建模为可计算、可检查的规则条目（如标题层级、摘要要素、关键词格式、图表标题与引用、参考文献格式、页边距/字体/字号/行距、中英文混排、三线表、分页与章节起始等），并设计规范交付质量诊断指标（见图4）。', indent=True)
add_figure_placeholder(c, '图4　校本规则库构建流程图', '图注：校本规则库由“校规—经验—错误”三源融合、迭代修正而成。')
add_para(c, '研究内容三：本科论文规范交付智能助手优化与实现', bold=True)
add_para(c, '在已有系统基础上，完善提交前体检、学术表达优化、格式对齐与标准 Word 导出四个环节，强化文档解析稳定性、规则匹配准确性与人工确认机制，使系统由“通用在线论文辅助工具”升级为“校本化规范交付智能助手”。', indent=True)
add_figure_placeholder(c, '图3　论文规范交付闭环流程图', '图注：系统形成“先发现问题、再人机协同修改、最后规范导出”的闭环。')
add_para(c, '研究内容四：系统应用验证、效果评价与推广模式总结', bold=True)
add_para(c, '设计应用验证方案与评价指标，采集试运行数据、开展问卷与访谈，分析系统在降低修改成本、提升提交质量、辅助教师指导等方面的成效，凝练校本化推广方案。', indent=True)
add_para(c, '工作方案：', bold=True)
add_para(c, '采用“文献与规范研究 + 系统迭代开发 + 校内试运行与教学应用 + 数据分析与评价”相结合的方法，遵循“需求调研→规范解析→规则库构建→系统模块完善→试运行数据采集→教学应用验证→评价指标分析→校本化推广方案形成”的技术路线（见图7）。', indent=True)
add_figure_placeholder(c, '图7　研究技术路线图', '图注：项目遵循“需求—规则—系统—验证—推广”的技术路线。')
add_para(c, '重点与难点：', bold=True)
add_para(c, '（1）规则化转化（难点）：如何将学校论文规范与教师经验中大量隐性、模糊、情境化的要求，转化为稳定、可计算、可检查的规则条目；（2）写作主体边界（重点）：如何在提供学术表达优化的同时，坚持以学生为写作主体，明确人机协同的合理边界，避免辅助工具替代学生自身写作；（3）文档解析稳定性（难点）：如何实现对不同来源、不同排版习惯论文文档的稳定解析、结构识别与规范化导出，保证内容不丢失；（4）应用效果评价（重点）：如何建立科学、可操作的应用效果评价指标体系；（5）小经费可验证（重点）：如何在校级项目经费较低的约束下，依托已有系统基础，完成可验证、可结题的阶段性成果。', indent=True)

# ================= ROW 5  四、现有研究基础（匿名）=================
c = cells[5]; clear_cell_keep_heading(c)
add_para(c, '本部分按匿名要求撰写，不涉及任何个人身份信息。本项目区别于一般“从零设想”的申报，其突出特点是已具备一套可演示、可运行、可继续迭代的系统基础。', indent=True)
add_para(c, '1. 已完成系统原型开发', bold=True)
add_para(c, '课题组已自主完成系统原型开发，系统具备论文上传、文档解析、提交前体检、学术表达优化、格式对齐与标准 Word 导出等功能，覆盖本科论文规范交付的主要环节，已可面向真实论文文档进行处理与演示（见图1）。', indent=True)
add_figure_placeholder(c, '图1　本科论文规范交付智能助手首页与核心功能界面',
    '图注：系统已形成“提交前体检—学术表达优化—格式对齐”三大功能入口，能够覆盖本科论文提交前的主要规范交付环节。占位说明：此处插入系统首页、模块入口与工具界面截图。', lines=6)
add_para(c, '2. 已完成云端部署', bold=True)
add_para(c, '系统已部署于阿里云服务器，采用前后端分离架构，具备在线访问、在线处理与持续迭代能力，能够支撑阶段性试运行与后续校内应用（系统技术架构见图5）。', indent=True)
add_figure_placeholder(c, '图5　本科论文规范交付智能助手技术架构图',
    '图注：系统由前端界面层、后端服务层、智能服务层、文档处理层、数据支撑层与部署层构成，已实现云端在线运行。', lines=5)
add_para(c, '3. 已形成可视化界面', bold=True)
add_para(c, '系统已形成首页、论文工具入口、提交前体检界面、学术表达优化界面、格式对齐界面与课程工作台入口等可视化界面，交互流程完整，后续申报与结题材料可直接插入相应截图。', indent=True)
add_para(c, '4. 已有阶段性运营基础', bold=True)
add_para(c, '系统已进行阶段性试运行，初步服务于本科论文修改与规范交付场景，积累了一定的访问、处理与反馈数据基础（具体数据见表1，由课题组后续以真实统计结果替换）。', indent=True)
add_table_block(c, '表1　系统阶段性试运行数据统计表',
    ['指标', '数据', '说明'],
    [['试运行时间', '【填写：如2025年X月至今，已持续X个月】', '自【年份/月】起持续迭代'],
     ['累计访问用户', '【填写：X 人】', '包括学生、教师与测试用户'],
     ['累计处理论文文档', '【填写：X 篇】', '包括格式对齐与体检任务'],
     ['累计表达优化请求', '【填写：X 次】', '反映学生对学术表达优化的需求'],
     ['累计导出标准 Word', '【填写：X 份】', '反映规范交付能力'],
     ['累计格式对齐任务', '【填写：X 次】', '反映格式规范化处理规模'],
     ['用户满意度', '【填写：X%】', '来自问卷或访谈反馈'],
     ['当前部署环境', '【阿里云服务器、Nginx、前后端分离、数据库/文件存储】', '支撑在线访问与持续迭代'],
     ['当前运营方式', '【试运行收费/邀请测试/小范围使用】', '说明存在真实使用需求'],
     ['获批后服务方式', '【面向呼伦贝尔学院师生免费开放/公益化试用】', '体现公益转化']])
add_para(c, '5. 已有用户反馈', bold=True)
add_para(c, '在试运行过程中，系统获得了学生与教师的初步认可，反馈相对集中于“减少格式修改时间”“提交前问题更清晰”“表达优化操作较直观”“标准 Word 导出较方便”等方面（反馈摘要见表2，均作匿名处理）。相关表述以稳健、客观为原则，不作夸大宣传。', indent=True)
add_table_block(c, '表2　用户反馈摘要表（匿名）',
    ['反馈维度', '典型表述（匿名处理）', '对应系统能力/拟改进方向'],
    [['修改效率', '【如：格式修改时间明显减少】', '格式对齐、标准导出'],
     ['问题定位', '【如：提交前知道要改哪里】', '提交前体检、问题清单'],
     ['表达优化', '【如：选中改写比较直观】', '学术表达优化、人机协同'],
     ['交付便利', '【如：导出 Word 可直接用】', '标准 Word 导出'],
     ['改进建议', '【如：希望覆盖更多格式细节】', '规则库迭代']],
    note='说明：反馈表述均需匿名，避免出现姓名、班级、学号等身份信息，填写时保持客观稳健。')
add_para(c, '6. 已有研究与技术基础', bold=True)
add_para(c, '课题组具备人工智能应用、软件开发、文档解析、Web 系统开发、教育教学实践与本科论文指导等方面的工作积累，能够支撑规则建模、系统迭代、文档处理与应用验证等研究任务。', indent=True)
add_para(c, '7. 校级项目转化基础', bold=True)
add_para(c, '当前系统处于小范围试运行/半商业化阶段，这一事实从侧面说明其具有真实使用需求与应用价值。若本项目获批校级科研项目，系统将进一步去商业化、公益化，面向呼伦贝尔学院师生免费开放或公益化试用，转化为服务学校本科人才培养与毕业论文指导的教育数字化工具。', indent=True)
add_para(c, '综上，项目具备低成本、短周期、可验证、可推广的实施基础：研究不依赖大额经费投入，而是以已有系统为支点，重点完成“通用工具→校本化智能助手”的升级与校内应用验证。', indent=True)

# ================= ROW 6  五、研究预计达到的目标 =================
c = cells[6]; clear_cell_keep_heading(c)
add_para(c, '1. 构建一套《呼伦贝尔学院本科论文规范交付规则清单》，将学校规范、教师经验与学生常见错误转化为可计算、可检查的规则条目；', indent=True)
add_para(c, '2. 完善一套本科论文规范交付智能助手系统，实现提交前体检、学术表达优化、格式对齐与标准 Word 导出的稳定运行；', indent=True)
add_para(c, '3. 形成一套“提交前体检—学术表达优化—格式对齐—标准导出”的规范交付闭环流程，并明确人机协同的合理边界；', indent=True)
add_para(c, '4. 完成系统校内试运行与教学应用评价，形成论文、应用报告与校本化推广方案，验证项目的应用成效与推广可行性。', indent=True)

# ================= ROW 7  六、计划研究进度 =================
c = cells[7]; clear_cell_keep_heading(c)
add_para(c, '按两年周期分四个阶段实施（见表3）。', indent=True)
add_table_block(c, '表3　项目研究进度安排表',
    ['阶段', '时间', '主要任务', '阶段成果'],
    [['第一阶段', '2026.06—2026.12', '需求调研、学校规范梳理、教师经验与学生常见错误收集、规则库初建', '需求分析报告、规则清单初稿'],
     ['第二阶段', '2027.01—2027.06', '系统模块完善、架构优化、规则匹配与文档解析增强、试运行数据采集', '系统优化版本、阶段性试运行数据'],
     ['第三阶段', '2027.07—2027.12', '教学应用验证、问卷与访谈、效果分析', '应用测试报告、论文初稿'],
     ['第四阶段', '2028.01—2028.06', '成果凝练、论文发表、规则清单定稿、结题材料整理', '论文、系统、规则清单、结题报告']])

# ================= ROW 8  七、项目成果形式 =================
c = cells[8]; clear_cell_keep_heading(c)
for line in [
    '1. 知网可检索论文 1 篇；',
    '2. 《呼伦贝尔学院本科论文规范交付规则清单》1 份；',
    '3. 本科论文规范交付智能助手系统 1 套（在线可演示）；',
    '4. 系统应用测试报告 1 份；',
    '5. 用户反馈与效果评价报告 1 份（匿名）；',
    '6. 本科论文规范交付流程图与校本化推广方案 1 份；',
    '7. 系统界面截图、技术架构图、功能说明文档若干。',
]:
    add_para(c, line, indent=True)
add_table_block(c, '表4　项目成果形式矩阵表',
    ['成果类型', '成果名称', '形式', '验收依据'],
    [['论文成果', '本科论文规范交付相关论文', '知网可检索论文 1 篇', '发表证明'],
     ['系统成果', '本科论文规范交付智能助手', '原型系统/在线平台', '系统截图、演示地址、功能说明'],
     ['规则成果', '呼伦贝尔学院本科论文规范交付规则清单', '文档', '规则清单'],
     ['应用成果', '系统试运行与教学应用报告', '报告', '数据统计、问卷反馈'],
     ['推广成果', '校本化论文规范交付流程方案', '方案', '流程图、制度建议']])

# ================= 参考文献（接在第七节单元格末尾）=================
def add_ref(cell, text):
    p = cell.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(18)
    pf.first_line_indent = Pt(-18)  # 悬挂缩进
    pf.line_spacing = 1.15
    pf.space_after = Pt(2)
    r = p.add_run(text)
    style_run(r, size=10.5)

c = cells[8]
add_para(c, '参考文献', bold=True, align=C, size=12, space_before=10)
REFS = [
 '[1] 教育数字化转型：转什么，怎么转[J]. 华东师范大学学报(教育科学版), 2023, 41(3): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.03.001.',
 '[2] 教育数字化转型的关键路径[J]. 华东师范大学学报(教育科学版), 2023, 41(3): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.03.007.',
 '[3] 胡姣, 彭红超, 祝智庭. 教育数字化转型的现实困境与突破路径[J]. 【刊名】, 2023, 【卷(期)】: 【页码】.',
 '[4] 杨宗凯, 王俊, 吴砥, 陈旭. ChatGPT/生成式人工智能对教育的影响探析及应对策略[J]. 华东师范大学学报(教育科学版), 2023, 41(7): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.07.003.',
 '[5] 朱永新, 杨帆. ChatGPT/生成式人工智能与教育创新：机遇、挑战以及未来[J]. 华东师范大学学报(教育科学版), 2023, 41(7): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.07.001.',
 '[6] 荀渊. ChatGPT/生成式人工智能与高等教育的价值和使命[J]. 华东师范大学学报(教育科学版), 2023, 41(7): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.07.006.',
 '[7] 陈静远, 胡丽雅, 吴飞. ChatGPT/生成式人工智能促进以知识点为核心的教学模式变革研究[J]. 华东师范大学学报(教育科学版), 2023, 41(7): 【页码】. DOI:10.16382/j.cnki.1000-5560.2023.07.016.',
 '[8] OO. The Impact of Automated Writing Evaluation on Second Language Writing Skills of Chinese EFL Learners: A Randomized Controlled Trial[J]. Frontiers in Psychology, 2023, 14. DOI:10.3389/fpsyg.2023.1249991.（著者待补）',
 '[9] OO. Automated Feedback and Writing: A Multi-level Meta-analysis of Effects on Students’ Performance[J]. 【刊名】, 2023. PMCID:PMC10351274.（著者待补）',
 '[10] Zhai N, Ma X. The Effectiveness of Automated Writing Evaluation on Writing Quality: A Meta-Analysis[J]. Journal of Educational Computing Research, 2023, 61(4). DOI:10.1177/07356331221127300.',
 '[11] OO. The Impact of Generative AI on Academic Reading and Writing: A Synthesis of Recent Evidence (2023–2025)[J]. Frontiers in Education, 2025. DOI:10.3389/feduc.2025.1711718.（著者待补）',
 '[12] OO. Can Large Language Models Provide Feedback to Students? A Case Study on ChatGPT[J/OL]. 2023.（刊名、著者待补；可经标题检索）',
 '[13] OO. The Impact of ChatGPT Feedback on the Development of EFL Students’ Writing Skills[J]. Cogent Education, 2024, 11(1). DOI:10.1080/2331186X.2024.2410101.（著者待补）',
 '[14] OO. Large Language Models Are State-of-the-Art Evaluators for Grammatical Error Correction[C]//Proceedings of the 19th Workshop on Innovative Use of NLP for Building Educational Applications (BEA). ACL, 2024. ACL Anthology: 2024.bea-1.6.（著者待补）',
 '[15] OO. Rethinking the Roles of Large Language Models in Chinese Grammatical Error Correction[J/OL]. arXiv preprint, 2024. arXiv:2402.11420.（著者待补）',
 '[16] OO. 生成式AI赋能研究生科研写作的学术伦理与风险防控[J]. 【刊名】, 2024, 【卷(期)】: 【页码】.（著者待补）',
 '[17] 廖祖君, 卜美文. 生成式人工智能时代学术生态的变革重塑、伦理争端与前瞻应对——基于期刊的视角[J]. 西华大学学报(哲学社会科学版), 2025, 【卷(期)】: 【页码】. DOI:10.12189/j.issn.1672-8505.2025.06.001.',
 '[18] 苗逢春. 生成式人工智能及其教育应用的基本争议和对策[J]. 开放教育研究, 2024, 【卷(期)】: 【页码】.',
 '[19] 王雯, 李永智. 国际生成式人工智能教育应用与省思[J]. 开放教育研究, 2024, 【卷(期)】: 【页码】.',
 '[20] 《学位法》框架下学位论文撰写中生成式人工智能应用的合规性研究[J]. 学位与研究生教育, 2024, 【卷(期)】: 【页码】.（著者待补）',
 '[21] OO. Human-AI Collaboration Patterns in AI-assisted Academic Writing[J]. Studies in Higher Education, 2024. DOI:10.1080/03075079.2024.2323593.（著者待补）',
 '[22] OO. Human-AI Collaboration in Academic Writing: Exploring University Students’ Agency Through a Sociocultural Lens[J]. 【刊名待核】, 2025. DOI/来源:ScienceDirect S2215039026000032.（著者待补）',
 '[23] OO. 【本科毕业论文质量相关论文，题名待补】[J]. 中国大学教学, 2022, (3): 【页码】.',
 '[24] 戴卫东. 提高本科生毕业论文质量的对策建议[J]. 【刊名】, 【年(期)】: 【页码】.',
 '[25] OO. “科技论文写作”课程教学研究现状分析与优化路径[J]. 中国林业教育, 2023. DOI:10.3969/j.1001-7232.20230027.（著者待补）',
 '[26] OO. 本科毕业论文外审和选题质量提升的教学改革研究[J]. 教育进展, 2024, 14(9): 【页码】.（著者待补）',
 '[27] OO. Generative AI: Is It a Paradigm Shift for Higher Education?[J]. Studies in Higher Education, 2024, 49(5). DOI:10.1080/03075079.2024.2332944.（著者待补）',
 '[28] OO. Generative AI and Higher Education: A Review of Claims from the First Months of ChatGPT[J]. Higher Education, 2024. DOI:10.1007/s10734-024-01265-3.（著者待补）',
]
for ref in REFS:
    add_ref(c, ref)

import os
os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print('SAVED', OUT)
print('paras row2:', len(cells[2].paragraphs), 'tables in doc:', len(doc.tables))
