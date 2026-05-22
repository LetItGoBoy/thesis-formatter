"""
.docx 格式化核心
backend/lib/formatter.py

按 config/formats/<template>.json 的规则「重建」一个新 .docx（不硬编码任何数值）。

实现要点：
- disable_snap_to_grid：所有段落格式化前先取消文档网格
- 中英文混排字体（字符级）：中文按段落字体，数字/英文/英文标点统一 ascii_font（TNR）
- 五个大块各另起一页；正文大块内每个 h1 另起一页
- 关键词清理为空格分隔；数字序号 1、->1.；章号双全角空格；固定标题文字
- 三线表：cantSplit + 上下 1.5pt、表头下 1pt（对文档中存在的表格生效）
"""
import io
import re
import base64
from copy import deepcopy

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shape import CT_Inline
from docx.opc.part import Part
from docx.opc.packuri import PackURI
from docx.image.exceptions import UnrecognizedImageError

from .docx_images import extract_ordered_images

# 图片关系类型 URI（OPC 规范）
_IMG_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"
# 不被 add_picture 支持的格式的扩展名映射
_EXT_BY_CT = {
    "image/x-emf": "emf", "image/emf": "emf",
    "image/x-wmf": "wmf", "image/wmf": "wmf",
    "image/png": "png",   "image/jpeg": "jpg",
    "image/gif": "gif",   "image/bmp": "bmp",
    "image/tiff": "tif",
}

# 中文 + 中文标点
CHINESE_RE = re.compile(r"[　-〿一-鿿＀-￯]")
FULL_WIDTH_SPACE = "　"

ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_BLOCK_ORDER = ["toc", "abstract", "body", "conclusion", "references"]


# ============================================================
# 取消文档网格
# ============================================================
def disable_snap_to_grid(paragraph):
    """取消单段对文档网格的对齐（设置字体/行距前必须先调用）"""
    pPr = paragraph._p.get_or_add_pPr()
    snap = pPr.find(qn("w:snapToGrid"))
    if snap is None:
        snap = OxmlElement("w:snapToGrid")
        pPr.append(snap)
    snap.set(qn("w:val"), "0")


def disable_doc_grid(doc):
    """取消 section 级文档网格"""
    for section in doc.sections:
        sectPr = section._sectPr
        grid = sectPr.find(qn("w:docGrid"))
        if grid is None:
            grid = OxmlElement("w:docGrid")
            sectPr.append(grid)
        grid.set(qn("w:type"), "default")
        grid.set(qn("w:linePitch"), "312")
        grid.set(qn("w:charSpace"), "0")


# ============================================================
# run 属性
# ============================================================
def _get_or_create_rpr(r_el):
    rpr = r_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        r_el.insert(0, rpr)
    return rpr


def _set_rpr_font(rpr, chinese_font, ascii_font):
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rfonts.set(qn("w:eastAsia"), chinese_font)


def _set_rpr_size(rpr, pt):
    half = str(int(float(pt) * 2))
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), half)


def _set_para_mark_size(paragraph, pt):
    """设置段落标记（pilcrow）的字号，决定空段落的可见高度。"""
    pPr = paragraph._p.get_or_add_pPr()
    rpr = pPr.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        pPr.append(rpr)
    _set_rpr_size(rpr, pt)


def _make_blank_line(paragraph, style):
    """构造一个真实的空行段落，高度 = 字号 × 行距，用于「空一行」分隔。"""
    disable_snap_to_grid(paragraph)
    pf = paragraph.paragraph_format
    ls = float(style.get("line_spacing", 1.5))
    pf.line_spacing = ls
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _set_para_mark_size(paragraph, float(style.get("font_size_pt", 12)))


def _set_rpr_bold(rpr, bold):
    b = rpr.find(qn("w:b"))
    bcs = rpr.find(qn("w:bCs"))
    if bold:
        if b is None:
            b = OxmlElement("w:b"); rpr.append(b)
        b.set(qn("w:val"), "1")
        if bcs is None:
            bcs = OxmlElement("w:bCs"); rpr.append(bcs)
        bcs.set(qn("w:val"), "1")
    else:
        if b is not None: rpr.remove(b)
        if bcs is not None: rpr.remove(bcs)


# ============================================================
# 中英文字符级拆分
# ============================================================
def _is_chinese(ch):
    return bool(CHINESE_RE.match(ch))


def _split_by_script(text):
    if not text:
        return []
    out, cur = [], [text[0]]
    cur_s = "zh" if _is_chinese(text[0]) else "en"
    for ch in text[1:]:
        s = "zh" if _is_chinese(ch) else "en"
        if s == cur_s:
            cur.append(ch)
        else:
            out.append(("".join(cur), cur_s))
            cur, cur_s = [ch], s
    out.append(("".join(cur), cur_s))
    return out


def _apply_mixed_font(paragraph, chinese_font, ascii_font, pt, bold):
    runs = list(paragraph.runs)
    if not runs:
        return
    for run in runs:
        text = run.text
        r_el = run._r
        if not text:
            rpr = _get_or_create_rpr(r_el)
            _set_rpr_font(rpr, chinese_font, ascii_font)
            _set_rpr_size(rpr, pt)
            _set_rpr_bold(rpr, bold)
            continue
        segments = _split_by_script(text)
        if len(segments) == 1:
            rpr = _get_or_create_rpr(r_el)
            _set_rpr_font(rpr, chinese_font, ascii_font)
            _set_rpr_size(rpr, pt)
            _set_rpr_bold(rpr, bold)
            continue
        orig_rpr = r_el.find(qn("w:rPr"))
        new_rs = []
        for seg, _ in segments:
            nr = OxmlElement("w:r")
            nrpr = deepcopy(orig_rpr) if orig_rpr is not None else OxmlElement("w:rPr")
            _set_rpr_font(nrpr, chinese_font, ascii_font)
            _set_rpr_size(nrpr, pt)
            _set_rpr_bold(nrpr, bold)
            nr.append(nrpr)
            t = OxmlElement("w:t")
            t.text = seg
            if seg != seg.strip() or " " in seg:
                t.set(qn("xml:space"), "preserve")
            nr.append(t)
            new_rs.append(nr)
        parent = r_el.getparent()
        idx = list(parent).index(r_el)
        for off, nr in enumerate(new_rs):
            parent.insert(idx + off, nr)
        parent.remove(r_el)


def _approx_width_pt(text: str, pt: float) -> float:
    """Approximate rendered width of text in pt.
    CJK / full-width chars ≈ pt wide; half-width space ≈ 0.5 pt; other ASCII ≈ 0.55 pt.
    Used for computing instructor left-indent to align with centered author line."""
    w = 0.0
    for ch in text:
        if CHINESE_RE.match(ch) or ch == "　":
            w += pt
        elif ch == " ":
            w += pt * 0.5
        else:
            w += pt * 0.55
    return w


# ============================================================
# 文本归一化
# ============================================================
# 关键词类型 → 前缀与是否首字母大写（config 未显式标记 is_keywords，按类型识别）
_KEYWORDS_META = {
    "keywords_cn": {"prefix": "关键词：", "capitalize": False},
    "keywords_en": {"prefix": "Keywords:", "capitalize": True},
}


def _normalize_text(text, ptype, style):
    if text is None:
        return ""

    kw = _KEYWORDS_META.get(ptype)
    if kw or style.get("is_keywords"):
        prefix = kw["prefix"] if kw else style.get("keywords_prefix", "")
        capitalize = kw["capitalize"] if kw else False
        body = text
        for p in [prefix, prefix.rstrip(":：") + ":", prefix.rstrip(":：") + "："]:
            if p and body.startswith(p):
                body = body[len(p):]
                break
        body = body.strip().rstrip("。.")
        if capitalize:
            # 英文：仅按标点切分（保留词内空格如"smart cane"），每个关键词首字母大写
            parts = [w.strip() for w in re.split(r"[;；,，、]+", body) if w.strip()]
            parts = [w[0].upper() + w[1:] if w else w for w in parts]
        else:
            # 中文：按标点和空白切分
            parts = [w for w in re.split(r"[;；,，、\s]+", body) if w]
        joined = FULL_WIDTH_SPACE.join(parts)
        return f"{prefix}{FULL_WIDTH_SPACE}{joined}" if prefix else joined

    # 作者行：将"作者"补全角空格，使其与"指导教师"等宽，居中时视觉对齐
    if ptype == "author_line":
        text = re.sub(r"^(作)\s*(者)", r"作　　者", text)
        return text

    if style.get("normalize_numbered_prefix"):
        text = re.sub(r"^(\s*\d+)\s*[、,，]\s*", r"\1. ", text)

    # 支持 chapter_spacing（int）或旧字段 chapter_two_space_normalize（bool）
    chapter_sp = int(style.get("chapter_spacing", 0)) or (
        2 if style.get("chapter_two_space_normalize") else 0
    )
    if chapter_sp > 0:
        m = re.match(r"^(第[一二三四五六七八九十百零\d]+章)\s*(.*)$", text.strip())
        if m and m.group(2):
            text = f"{m.group(1)}{FULL_WIDTH_SPACE * chapter_sp}{m.group(2)}"

    # 支持 fixed_content（新字段）或旧字段 fixed_text
    fixed = style.get("fixed_content") or style.get("fixed_text")
    if fixed:
        bare = text.replace(" ", "").replace(FULL_WIDTH_SPACE, "")
        if bare == fixed.replace(" ", "").replace(FULL_WIDTH_SPACE, ""):
            text = fixed

    n_lead = int(style.get("leading_full_width_spaces", 0))
    if n_lead > 0:
        stripped = text.lstrip(FULL_WIDTH_SPACE).lstrip()
        text = FULL_WIDTH_SPACE * n_lead + stripped

    return text


# ============================================================
# 段落样式
# ============================================================
def _apply_paragraph_style(paragraph, ptype, style, skip_page_break=False):
    disable_snap_to_grid(paragraph)  # 先取消网格

    pf = paragraph.paragraph_format
    align = style.get("alignment", "justify")
    if align in ALIGNMENT_MAP:
        paragraph.alignment = ALIGNMENT_MAP[align]

    ls = style.get("line_spacing")
    if ls is not None:
        pf.line_spacing = float(ls)
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    # space_before_lines → 转为 pt（行数 × 字号 × 行距倍数）
    sbl = float(style.get("space_before_lines", 0))
    if sbl > 0:
        _pt = float(style.get("font_size_pt", 12))
        _ls = float(style.get("line_spacing", 1.5))
        pf.space_before = Pt(sbl * _pt * _ls)
    else:
        pf.space_before = Pt(float(style.get("space_before_pt", 0)))
    pf.space_after = Pt(float(style.get("space_after_pt", 0)))

    pt = float(style.get("font_size_pt", 12))
    first_chars = style.get("first_line_indent_chars", 0)
    hanging     = style.get("hanging_indent_chars", 0)
    left_chars  = style.get("indent_chars", 0)          # 整段左缩进（目录二/三级）
    if hanging and hanging > 0:
        pf.left_indent       = Pt(hanging * pt)
        pf.first_line_indent = Pt(-hanging * pt)
    elif left_chars and int(left_chars) > 0:
        pf.left_indent       = Pt(int(left_chars) * pt)
        pf.first_line_indent = Pt(0)                    # 显式清零，防继承
    elif first_chars and first_chars > 0:
        pf.first_line_indent = Pt(first_chars * pt)
        pf.left_indent       = None
    else:
        pf.first_line_indent = None
        pf.left_indent       = None

    if style.get("page_break_before") and not skip_page_break:
        pf.page_break_before = True

    # 支持 font_cn（新字段）或旧字段 chinese_font；同理 font_en_num / ascii_font
    cn_font = style.get("font_cn") or style.get("chinese_font", "宋体")
    en_font = style.get("font_en_num") or style.get("ascii_font", "Times New Roman")
    _apply_mixed_font(paragraph, cn_font, en_font, pt, bool(style.get("bold", False)))


# ============================================================
# ============================================================
# 段落大纲级别（导航窗口 / 文档结构图）
# ============================================================
_OUTLINE_LEVEL: dict[str, int] = {
    "toc_title": 0,
    "h1": 0, "h2": 1, "h3": 2,
    "conclusion_title": 0,
    "references_title": 0,
}


def _set_outline_level(paragraph, level: int):
    """设置 w:outlineLvl，让 WPS/Word 导航窗口识别标题层次。"""
    pPr = paragraph._p.get_or_add_pPr()
    old = pPr.find(qn("w:outlineLvl"))
    if old is not None:
        pPr.remove(old)
    el = OxmlElement("w:outlineLvl")
    el.set(qn("w:val"), str(level))
    pPr.append(el)


# ============================================================
# 三线表（对文档中存在的表格生效）
# ============================================================
def _set_cell_borders(cell, borders):
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    tcb = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        v = borders.get(edge, 0)
        if v and v > 0:
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), str(int(v * 8)))
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "000000")
        else:
            e.set(qn("w:val"), "nil")
        tcb.append(e)
    tcPr.append(tcb)


def _embed_image_raw(doc, run_el, blob, content_type, cx_emu, cy_emu, max_width_emu):
    """直接把 blob 注册进文档包，绕过 add_picture 的格式校验（用于 EMF/WMF 等矢量图）。"""
    if not cx_emu or cx_emu <= 0:
        cx_emu = max_width_emu
    if not cy_emu or cy_emu <= 0:
        cy_emu = int(cx_emu * 3 / 4)
    cx_emu = min(int(cx_emu), max_width_emu)

    ext = _EXT_BY_CT.get(content_type, "bin")
    # 用当前包内部件数量做唯一编号，避免 partname 冲突
    n = sum(1 for _ in doc.part.package.iter_parts())
    partname = PackURI(f"/word/media/img_{n}.{ext}")

    img_part = Part(partname, content_type, blob, doc.part.package)
    rId = doc.part.relate_to(img_part, _IMG_REL)

    shape_id = doc.part.next_id
    inline = CT_Inline.new_pic_inline(shape_id, rId, f"image{shape_id}.{ext}", cx_emu, cy_emu)
    drawing = OxmlElement("w:drawing")
    drawing.append(inline)
    run_el.append(drawing)


def _add_image_paragraph(doc, p, images, max_width_emu, need_page_break):
    """居中插入一张图片段落。优先用原文 image_index 取图，缺失时回退前端 image_b64。"""
    para = doc.add_paragraph()
    if need_page_break:
        para.paragraph_format.page_break_before = True
    disable_snap_to_grid(para)
    para.alignment = ALIGNMENT_MAP["center"]
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)

    blob = None
    content_type = "image/png"
    cx = cy = None
    img = None
    ii = p.get("image_index")
    if isinstance(ii, int) and 0 <= ii < len(images):
        img = images[ii]
    if img:
        blob = img.get("blob")
        cx = img.get("cx")
        cy = img.get("cy")
        content_type = img.get("content_type") or content_type
    if blob is None:  # 源文件取图失败，回退前端回传的 base64（若有）
        b64 = p.get("image_b64") or ""
        if "," in b64:
            try:
                # 从 data URI 解析 content_type
                header, data = b64.split(",", 1)
                if ";" in header:
                    content_type = header.split(":")[1].split(";")[0]
                blob = base64.b64decode(data)
            except Exception:
                blob = None

    if not blob:
        para.add_run("[图片缺失]")
        return para

    width_emu = Emu(min(int(cx), max_width_emu)) if cx and cx > 0 else Emu(max_width_emu)
    run = para.add_run()
    try:
        run.add_picture(io.BytesIO(blob), width=width_emu)
    except UnrecognizedImageError:
        # EMF/WMF 等矢量图：绕过格式校验直接嵌入，Word/WPS 可原生渲染
        _embed_image_raw(doc, run._r, blob, content_type, cx, cy, max_width_emu)
    except Exception:
        para.add_run("[图片无法载入]")
    return para


def _add_table(doc, cells):
    """按二维文本网格新建表格。样式（三线/字体/居中）由 _format_tables 之后统一施加。"""
    rows = len(cells)
    cols = max((len(r) for r in cells), default=0)
    if rows == 0 or cols == 0:
        return None
    table = doc.add_table(rows=rows, cols=cols)
    for i, row in enumerate(cells):
        for j in range(cols):
            table.rows[i].cells[j].text = row[j] if j < len(row) else ""
    return table


def _format_tables(doc, table_cfg):
    cant_split = bool(table_cfg.get("cant_split", True))
    tl = table_cfg.get("three_line", {})
    on = bool(tl.get("enabled", True))
    top = float(tl.get("top_border_pt", 1.5))
    bottom = float(tl.get("bottom_border_pt", 1.5))
    header = float(tl.get("header_bottom_border_pt", 1.0))
    inner = float(tl.get("inner_border_pt", 0))

    tfont = table_cfg.get("font", {})
    zh = tfont.get("chinese_font", "宋体")
    en = tfont.get("ascii_font", "Times New Roman")
    sz = float(tfont.get("font_size_pt", 10.5))

    for table in doc.tables:
        rows = table.rows
        n = len(rows)
        if n == 0:
            continue
        if cant_split:
            for row in rows:
                trPr = row._tr.get_or_add_trPr()
                if trPr.find(qn("w:cantSplit")) is None:
                    trPr.append(OxmlElement("w:cantSplit"))
        if on:
            for i, row in enumerate(rows):
                for cell in row.cells:
                    if n == 1:
                        b = {"top": top, "bottom": bottom}
                    elif i == 0:
                        b = {"top": top, "bottom": header}
                    elif i == n - 1:
                        b = {"top": inner, "bottom": bottom}
                    else:
                        b = {"top": inner, "bottom": inner}
                    _set_cell_borders(cell, b)
        for row in rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for para in cell.paragraphs:
                    disable_snap_to_grid(para)
                    para.alignment = ALIGNMENT_MAP["center"]
                    _apply_mixed_font(para, zh, en, sz, False)


# ============================================================
# 页面设置
# ============================================================
def _apply_page_settings(doc, page_cfg):
    for section in doc.sections:
        if "margin_top_cm" in page_cfg:
            section.top_margin = Cm(float(page_cfg["margin_top_cm"]))
        if "margin_bottom_cm" in page_cfg:
            section.bottom_margin = Cm(float(page_cfg["margin_bottom_cm"]))
        if "margin_left_cm" in page_cfg:
            section.left_margin = Cm(float(page_cfg["margin_left_cm"]))
        if "margin_right_cm" in page_cfg:
            section.right_margin = Cm(float(page_cfg["margin_right_cm"]))
        sectPr = section._sectPr
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is not None:
            pgMar.set(qn("w:gutter"), str(int(float(page_cfg.get("gutter_cm", 0)) * 567)))


def _block_of(ptype, styles):
    blk = (styles.get(ptype) or {}).get("block", "body")
    return "body" if blk == "_legacy" else blk


# ============================================================
# 目录点状前导线
# ============================================================
_TOC_ENTRY_RE = re.compile(r'^(.+?)[\s.·]+(\d{1,4})\s*$')


def _split_toc_entry(text: str):
    """分离目录条目的标题与页码。无页码时返回 (text, None)。"""
    m = _TOC_ENTRY_RE.match(text.strip())
    if m and m.group(1).strip():
        return m.group(1).strip(), m.group(2)
    return text, None


def _add_dot_leader_tab(paragraph, text_width_twips: int):
    """在段落上设置右对齐+点状前导符制表位（位置=正文宽度，单位：twip）。"""
    pPr = paragraph._p.get_or_add_pPr()
    old = pPr.find(qn("w:tabs"))
    if old is not None:
        pPr.remove(old)
    tabs_el = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:leader"), "dot")
    tab.set(qn("w:pos"), str(text_width_twips))
    tabs_el.append(tab)
    pPr.append(tabs_el)


# ============================================================
# 主入口：重建 .docx
# ============================================================
def format_docx(paragraphs, format_config, source_bytes=None):
    """
    paragraphs: [{"index": int, "type": str, "text": str}, ...]
    format_config: hulunbeier_univ.json 内容
    source_bytes: 兼容参数；本实现按 CLAUDE.md「重建.docx」从段落列表新建文档。
    """
    doc = Document()
    styles = format_config.get("paragraph_styles", {})

    _apply_page_settings(doc, format_config.get("page", {}))
    disable_doc_grid(doc)

    # 正文宽度（twip）= (A4宽 - 左边距 - 右边距) × 567 twip/cm
    _page = format_config.get("page", {})
    _text_w_cm = (
        21.0
        - float(_page.get("margin_left_cm",  3.0))
        - float(_page.get("margin_right_cm", 2.0))
    )
    _text_w_twips = int(_text_w_cm * 567)
    _max_img_emu = int(_text_w_cm * 360000)  # 图片最大宽度 = 正文宽度（1cm=360000EMU）

    # 原文图片按文档顺序列表，下标 = 解析端分配的 image_index
    try:
        _images = extract_ordered_images(source_bytes) if source_bytes else []
    except Exception:
        _images = []

    ordered = sorted(paragraphs, key=lambda p: p.get("index", 0))

    # 参考文献自动补序号：若段落文字本身不带 [n] 前缀（Word 自动编号情形），自动注入
    _REF_TYPES = {"reference_item", "ref"}
    _ref_items = [p for p in ordered if p.get("type") in _REF_TYPES]
    _already_num = sum(
        1 for p in _ref_items
        if re.match(r"^\s*\[\d+\]", p.get("text", ""))
    )
    _auto_num_refs = bool(_ref_items) and _already_num < len(_ref_items) / 2
    _ref_seq = 0

    # 作者/指导老师首字对齐：缓存 author_line 居中后的左起点，供 instructor 使用
    _author_left_pt: float | None = None
    _text_w_pt = _text_w_cm * 28.35  # cm → pt

    last_block = None
    for p in ordered:
        ptype = p.get("type", "body")
        style = styles.get(ptype) or styles.get("body", {})
        cur_block = _block_of(ptype, styles)
        need_page_break = last_block is not None and cur_block != last_block
        last_block = cur_block
        # 是否需要另起一页：大块切换 或 段落样式自带 page_break_before
        starts_new_page = need_page_break or bool(style.get("page_break_before"))

        # 表格：用 cells 网格重建为真正的 Word 表格（三线表样式由 _format_tables 统一施加）
        if ptype == "table" and p.get("cells"):
            if need_page_break:
                brk = doc.add_paragraph()
                brk.paragraph_format.page_break_before = True
            _add_table(doc, p["cells"])
            continue

        # 图片：从原文按 image_index 取出原图，居中插回（题注按文档顺序自然位于其下方）
        if ptype == "figure":
            _add_image_paragraph(doc, p, _images, _max_img_emu, need_page_break)
            continue

        text = _normalize_text(p.get("text", ""), ptype, style)

        # 参考文献序号：自动补 [n] 前缀
        if ptype in _REF_TYPES and _auto_num_refs:
            _ref_seq += 1
            text = f"[{_ref_seq}] {text}"

        # 作者/指导老师首字对齐
        # author_line 居中后缓存其左起点 = (正文宽 - 行宽) / 2
        if ptype == "author_line":
            _pt = float(style.get("font_size_pt", 12))
            _author_left_pt = (_text_w_pt - _approx_width_pt(text, _pt)) / 2.0

        # 前置真实空行（blank_lines_before）。若本段要另起一页，则把分页放到第一个
        # 空行上，使空行落在新页顶部、正文随后；否则空行就在当前位置之前。
        n_blank = int(style.get("blank_lines_before", 0))
        page_break_consumed = False
        for i in range(n_blank):
            blank = doc.add_paragraph()
            if i == 0 and starts_new_page:
                blank.paragraph_format.page_break_before = True
                page_break_consumed = True
            _make_blank_line(blank, style)

        para = doc.add_paragraph()

        # 大块切换 -> 另起一页（若分页已由前置空行承担则跳过）
        if need_page_break and not page_break_consumed:
            para.paragraph_format.page_break_before = True

        # 目录条目：右对齐点状前导制表位 + 分离页码
        if ptype in ("toc_h1", "toc_h2", "toc_h3"):
            title, pagenum = _split_toc_entry(text)
            if pagenum:
                _add_dot_leader_tab(para, _text_w_twips)
                para.add_run(title)
                tab_r = OxmlElement("w:r")
                tab_r.append(OxmlElement("w:tab"))
                para._p.append(tab_r)
                para.add_run(pagenum)
            else:
                para.add_run(text)
        else:
            para.add_run(text)

        _apply_paragraph_style(para, ptype, style, skip_page_break=page_break_consumed)
        if ptype in _OUTLINE_LEVEL:
            _set_outline_level(para, _OUTLINE_LEVEL[ptype])

        # instructor 左缩进 = author_line 居中后的左起点，首字与作者行对齐
        if ptype == "instructor" and _author_left_pt is not None and _author_left_pt > 0:
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.left_indent = Pt(_author_left_pt)

    _format_tables(doc, format_config.get("table_format", {}))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
