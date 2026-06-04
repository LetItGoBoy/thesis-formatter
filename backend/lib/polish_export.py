"""
论文表达润色 - 导出模块
backend/lib/polish_export.py

基于原始 docx，只替换用户确认修改过的段落文本，其余（表格、图片、公式、
未修改段落）原样保留。遍历顺序与 polish_parser.import_docx_as_blocks 完全一致，
以 index 对齐回写。

替换段落时尽量保留段落级样式（对齐、缩进、样式名）：
清空原 runs 后用第一个 run 的字体属性 add_run（无 run 时直接 add_run）。
"""
import io
import copy
import logging

from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

logger = logging.getLogger("thesis.polish.export")


def _iter_block_items(doc: DocxDocument):
    """与 polish_parser 完全一致的遍历顺序，保证 index 对齐。"""
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield DocxParagraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield DocxTable(child, doc)


def _replace_paragraph_text(para: DocxParagraph, new_text: str) -> None:
    """
    替换段落文字，尽量保留段落级样式与首个 run 的字符样式。
    不动 paragraph_format（对齐/缩进/样式名），只重建 run 内容。
    """
    runs = para.runs
    # 保留第一个 run 的 rPr（字体/字号/加粗等字符样式）作为模板
    template_rpr = None
    if runs and runs[0]._element.rPr is not None:
        template_rpr = copy.deepcopy(runs[0]._element.rPr)

    # 删除原有所有 run
    for run in list(runs):
        run._element.getparent().remove(run._element)

    new_run = para.add_run(new_text)
    if template_rpr is not None:
        # 用模板 rPr 覆盖新 run 的字符样式，保持字体一致
        new_rpr = new_run._element.get_or_add_rPr()
        new_run._element.replace(new_rpr, copy.deepcopy(template_rpr))


def export_polished_docx(file_bytes: bytes, blocks: list[dict]) -> io.BytesIO:
    """
    打开原始 docx，按顺序遍历段落与表格，只替换 kind="paragraph" 且 changed=true
    的段落文字；表格、图片、公式、未修改段落均不动。返回 BytesIO。
    """
    doc = DocxDocument(io.BytesIO(file_bytes))

    # index -> block（仅保留需要替换的段落）
    changes: dict[int, str] = {}
    for b in blocks or []:
        if (
            b.get("kind") == "paragraph"
            and b.get("changed")
            and isinstance(b.get("index"), int)
        ):
            changes[b["index"]] = b.get("text", "")

    replaced = 0
    idx = 0
    for blk in _iter_block_items(doc):
        if isinstance(blk, DocxTable):
            idx += 1
            continue
        # 段落
        if idx in changes:
            _replace_paragraph_text(blk, changes[idx])
            replaced += 1
        idx += 1

    logger.info("export_polished_docx: 遍历 %d 块，替换 %d 段（请求修改 %d 段）",
                idx, replaced, len(changes))

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
